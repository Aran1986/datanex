# Location: datanex/core/file_handler.py

import pandas as pd
import polars as pl
from pathlib import Path
from typing import Dict, Any, Tuple
import magic
from utils.logger import log
from models.file import FileType
import json
import xml.etree.ElementTree as ET
from PyPDF2 import PdfReader
from docx import Document

class FileHandler:
    """ماژول 1: دریافت و پردازش اولیه فایل"""
    
    SUPPORTED_TYPES = {
        'text/csv': FileType.CSV,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': FileType.EXCEL,
        'application/vnd.ms-excel': FileType.EXCEL,
        'application/json': FileType.JSON,
        'text/xml': FileType.XML,
        'application/xml': FileType.XML,
        'application/pdf': FileType.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
        'text/plain': FileType.TXT,
        'text/html': FileType.HTML,
        'application/sql': FileType.SQL,
    }
    
    async def detect_file_type(self, file_data: bytes) -> Tuple[FileType, str]:
        """تشخیص نوع فایل"""
        mime = magic.from_buffer(file_data, mime=True)
        file_type = self.SUPPORTED_TYPES.get(mime, FileType.UNKNOWN)
        log.info(f"Detected file type: {file_type} (MIME: {mime})")
        return file_type, mime
    
    async def extract_metadata(self, file_data: bytes, file_type: FileType) -> Dict[str, Any]:
        """استخراج متادیتا از فایل"""
        metadata = {
            'size': len(file_data),
            'type': file_type.value
        }
        
        try:
            if file_type == FileType.CSV:
                df = pd.read_csv(pd.io.common.BytesIO(file_data), nrows=5)
                metadata.update({
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.astype(str).to_dict(),
                    'sample_rows': df.head(3).to_dict('records')
                })
            
            elif file_type == FileType.EXCEL:
                df = pd.read_excel(pd.io.common.BytesIO(file_data), nrows=5)
                metadata.update({
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.astype(str).to_dict(),
                    'sample_rows': df.head(3).to_dict('records')
                })
            
            elif file_type == FileType.JSON:
                data = json.loads(file_data.decode('utf-8'))
                metadata.update({
                    'structure': 'array' if isinstance(data, list) else 'object',
                    'keys': list(data[0].keys()) if isinstance(data, list) and len(data) > 0 else list(data.keys()),
                    'record_count': len(data) if isinstance(data, list) else 1
                })
            
            elif file_type == FileType.PDF:
                pdf = PdfReader(pd.io.common.BytesIO(file_data))
                metadata.update({
                    'pages': len(pdf.pages),
                    'text_sample': pdf.pages[0].extract_text()[:500] if len(pdf.pages) > 0 else ""
                })
            
            elif file_type == FileType.DOCX:
                doc = Document(pd.io.common.BytesIO(file_data))
                metadata.update({
                    'paragraphs': len(doc.paragraphs),
                    'text_sample': doc.paragraphs[0].text[:500] if len(doc.paragraphs) > 0 else ""
                })
            
        except Exception as e:
            log.error(f"Error extracting metadata: {e}")
            metadata['error'] = str(e)
        
        return metadata
    
    async def load_data(self, file_data: bytes, file_type: FileType) -> pd.DataFrame:
        """بارگذاری داده به DataFrame"""
        try:
            if file_type == FileType.CSV:
                return pd.read_csv(pd.io.common.BytesIO(file_data))
            
            elif file_type == FileType.EXCEL:
                return pd.read_excel(pd.io.common.BytesIO(file_data))
            
            elif file_type == FileType.JSON:
                data = json.loads(file_data.decode('utf-8'))
                if isinstance(data, list):
                    return pd.DataFrame(data)
                else:
                    return pd.DataFrame([data])
            
            elif file_type == FileType.XML:
                root = ET.fromstring(file_data.decode('utf-8'))
                records = []
                for child in root:
                    record = {elem.tag: elem.text for elem in child}
                    records.append(record)
                return pd.DataFrame(records)
            
            elif file_type == FileType.TXT:
                lines = file_data.decode('utf-8').split('\n')
                return pd.DataFrame({'text': lines})
            
            else:
                raise ValueError(f"Unsupported file type for data loading: {file_type}")
                
        except Exception as e:
            log.error(f"Error loading data: {e}")
            raise
    
    async def get_statistics(self, df: pd.DataFrame) -> Dict[str, Any]:
        """آمار کلی از داده"""
        return {
            'row_count': len(df),
            'column_count': len(df.columns),
            'columns': df.columns.tolist(),
            'dtypes': df.dtypes.astype(str).to_dict(),
            'null_counts': df.isnull().sum().to_dict(),
            'memory_usage': df.memory_usage(deep=True).sum(),
            'numeric_summary': df.describe().to_dict() if len(df.select_dtypes(include='number').columns) > 0 else {}
        }

file_handler = FileHandler()